"""
Document Ingestion Module

Handles document loading, validation, and preprocessing for the OCR + NLP pipeline.
Supports multiple document formats and batch processing.
"""

import os
import mimetypes
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, Generator
from dataclasses import dataclass
from PIL import Image
import PyPDF2
from pdf2image import convert_from_path
from docx import Document
import yaml
from loguru import logger


@dataclass
class DocumentMetadata:
    """Metadata for processed documents"""
    file_path: str
    file_name: str
    file_size: int
    mime_type: str
    format: str
    pages: int = 1
    created_at: Optional[str] = None
    modified_at: Optional[str] = None


@dataclass
class ProcessedDocument:
    """Container for processed document data"""
    metadata: DocumentMetadata
    content: Union[List[Image.Image], str]  # Images for OCR or text for direct processing
    page_count: int
    processing_time: Optional[float] = None


class DocumentIngestionError(Exception):
    """Custom exception for document ingestion errors"""
    pass


class DocumentIngestor:
    """
    Handles document ingestion with support for multiple formats.
    Designed for scalability and batch processing.
    """
    
    def __init__(self, config_path: str = "config/config.yaml"):
        self.config = self._load_config(config_path)
        self.supported_formats = self.config["documents"]["supported_formats"]
        self.max_file_size = self.config["documents"]["max_file_size_mb"] * 1024 * 1024  # Convert to bytes
        
        logger.info(f"DocumentIngestor initialized with formats: {self.supported_formats}")
    
    def _load_config(self, config_path: str) -> Dict[str, Any]:
        """Load configuration from YAML file"""
        try:
            with open(config_path, 'r') as file:
                return yaml.safe_load(file)
        except Exception as e:
            logger.error(f"Failed to load config from {config_path}: {e}")
            # Return default config if file loading fails
            return {
                "documents": {
                    "supported_formats": ["pdf", "png", "jpg", "jpeg", "tiff", "docx"],
                    "max_file_size_mb": 50
                }
            }
    
    def validate_document(self, file_path: str) -> bool:
        """
        Validate document format and size
        
        Args:
            file_path: Path to the document file
            
        Returns:
            bool: True if document is valid, False otherwise
        """
        if not os.path.exists(file_path):
            logger.error(f"File does not exist: {file_path}")
            return False
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size:
            logger.error(f"File too large: {file_size} bytes (max: {self.max_file_size})")
            return False
        
        # Check file format
        file_ext = Path(file_path).suffix.lower().lstrip('.')
        if file_ext not in self.supported_formats:
            logger.error(f"Unsupported format: {file_ext}")
            return False
        
        return True
    
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Extract metadata from document"""
        file_path_obj = Path(file_path)
        stat = file_path_obj.stat()
        
        # Get MIME type
        mime_type, _ = mimetypes.guess_type(file_path)
        
        # Determine format
        format_ext = file_path_obj.suffix.lower().lstrip('.')
        
        # Count pages for PDF documents
        pages = 1
        if format_ext == 'pdf':
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    pages = len(pdf_reader.pages)
            except Exception as e:
                logger.warning(f"Could not count PDF pages: {e}")
        
        return DocumentMetadata(
            file_path=file_path,
            file_name=file_path_obj.name,
            file_size=stat.st_size,
            mime_type=mime_type or "application/octet-stream",
            format=format_ext,
            pages=pages,
            created_at=str(stat.st_ctime),
            modified_at=str(stat.st_mtime)
        )
    
    def process_pdf(self, file_path: str, metadata: DocumentMetadata) -> List[Image.Image]:
        """Convert PDF to images for OCR processing"""
        try:
            # Convert PDF pages to images
            images = convert_from_path(file_path, dpi=300)
            logger.info(f"Converted PDF to {len(images)} images")
            return images
        except Exception as e:
            raise DocumentIngestionError(f"Failed to process PDF {file_path}: {e}")
    
    def process_image(self, file_path: str, metadata: DocumentMetadata) -> List[Image.Image]:
        """Load and validate image file"""
        try:
            image = Image.open(file_path)
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            return [image]
        except Exception as e:
            raise DocumentIngestionError(f"Failed to process image {file_path}: {e}")
    
    def process_docx(self, file_path: str, metadata: DocumentMetadata) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            return "\n".join(text_content)
        except Exception as e:
            raise DocumentIngestionError(f"Failed to process DOCX {file_path}: {e}")
    
    def ingest_single_document(self, file_path: str) -> ProcessedDocument:
        """
        Ingest a single document
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ProcessedDocument: Processed document with content and metadata
        """
        import time
        start_time = time.time()
        
        # Validate document
        if not self.validate_document(file_path):
            raise DocumentIngestionError(f"Document validation failed for {file_path}")
        
        # Extract metadata
        metadata = self.extract_metadata(file_path)
        logger.info(f"Processing {metadata.format.upper()} document: {metadata.file_name}")
        
        # Process based on format
        if metadata.format == 'pdf':
            content = self.process_pdf(file_path, metadata)
            page_count = len(content)
        elif metadata.format in ['png', 'jpg', 'jpeg', 'tiff']:
            content = self.process_image(file_path, metadata)
            page_count = 1
        elif metadata.format == 'docx':
            content = self.process_docx(file_path, metadata)
            page_count = 1
        else:
            raise DocumentIngestionError(f"Unsupported format: {metadata.format}")
        
        processing_time = time.time() - start_time
        
        return ProcessedDocument(
            metadata=metadata,
            content=content,
            page_count=page_count,
            processing_time=processing_time
        )
    
    def ingest_batch(self, file_paths: List[str]) -> Generator[ProcessedDocument, None, None]:
        """
        Ingest multiple documents in batch
        
        Args:
            file_paths: List of file paths to process
            
        Yields:
            ProcessedDocument: Each processed document
        """
        logger.info(f"Starting batch ingestion of {len(file_paths)} documents")
        
        for i, file_path in enumerate(file_paths, 1):
            try:
                logger.info(f"Processing document {i}/{len(file_paths)}: {Path(file_path).name}")
                document = self.ingest_single_document(file_path)
                yield document
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
                continue
    
    def ingest_directory(self, directory_path: str, recursive: bool = True) -> Generator[ProcessedDocument, None, None]:
        """
        Ingest all supported documents from a directory
        
        Args:
            directory_path: Path to the directory containing documents
            recursive: Whether to search subdirectories
            
        Yields:
            ProcessedDocument: Each processed document
        """
        directory = Path(directory_path)
        
        if not directory.exists():
            raise DocumentIngestionError(f"Directory does not exist: {directory_path}")
        
        # Find all supported files
        pattern = "**/*" if recursive else "*"
        all_files = []
        
        for file_path in directory.glob(pattern):
            if file_path.is_file():
                file_ext = file_path.suffix.lower().lstrip('.')
                if file_ext in self.supported_formats:
                    all_files.append(str(file_path))
        
        logger.info(f"Found {len(all_files)} supported documents in {directory_path}")
        
        # Process files in batches
        yield from self.ingest_batch(all_files)


def create_sample_documents():
    """Create sample documents for testing"""
    input_dir = Path("data/input")
    input_dir.mkdir(exist_ok=True)
    
    # Create a simple text file that can be converted to PDF for testing
    sample_text = """
    Sample Document for Testing
    
    This is a sample document created for testing the IntelliDoc pipeline.
    
    Key Information:
    - Company: Acme Corporation
    - Date: January 15, 2024
    - Amount: $5,000.00
    - Contact: John Doe (john.doe@acme.com)
    
    This document contains various entities that should be extracted by the NLP pipeline.
    """
    
    with open(input_dir / "sample_document.txt", "w") as f:
        f.write(sample_text)
    
    logger.info("Created sample documents for testing")


if __name__ == "__main__":
    # Test the ingestion system
    ingestor = DocumentIngestor()
    
    # Create sample documents
    create_sample_documents()
    
    # Test directory ingestion
    try:
        for doc in ingestor.ingest_directory("data/input"):
            print(f"Processed: {doc.metadata.file_name} ({doc.metadata.format})")
            print(f"Size: {doc.metadata.file_size} bytes")
            print(f"Processing time: {doc.processing_time:.2f}s")
            print("-" * 40)
    except Exception as e:
        logger.error(f"Ingestion test failed: {e}")
